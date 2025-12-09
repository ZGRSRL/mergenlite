#!/usr/bin/env python3
"""
Document Processor Agent - MergenLite Core
PDF indirme, metin çıkarımı ve yapılandırma
SAM.gov dokümanları için optimize edilmiş
"""

import io
import os
import json
import requests
import logging
import zipfile
from typing import Dict, Any, Optional, List
from pathlib import Path
from datetime import datetime

logger = logging.getLogger(__name__)

try:
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("pdfplumber not available, PDF processing will be limited")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available, DOCX processing will be limited")

try:
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    logger.warning("pandas not available, Excel processing will be limited")

class DocumentProcessor:
    """Doküman işleme servisi"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'doc', 'txt', 'zip', 'xls', 'xlsx']
    
    def process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """Yüklenen dosyayı işle"""
        
        try:
            file_extension = Path(uploaded_file.name).suffix.lower().replace('.', '')
            
            if file_extension not in self.supported_formats:
                return {
                    'success': False,
                    'error': f'Desteklenmeyen dosya formatı: {file_extension}'
                }
            
            # Dosya içeriğini oku
            file_content = uploaded_file.read()
            
            if file_extension == 'pdf':
                return self._process_pdf(file_content, uploaded_file.name)
            elif file_extension in ['docx', 'doc']:
                return self._process_docx(file_content, uploaded_file.name)
            elif file_extension == 'txt':
                # TXT için bytes'ı string'e çevir
                try:
                    text_content = file_content.decode('utf-8')
                except UnicodeDecodeError:
                    text_content = file_content.decode('latin-1', errors='ignore')
                return self._process_txt_content(text_content, uploaded_file.name)
            elif file_extension == 'zip':
                return self._process_zip(file_content, uploaded_file.name)
            elif file_extension in ['xls', 'xlsx']:
                return self._process_excel(file_content, uploaded_file.name)
            else:
                return {
                    'success': False,
                    'error': f'Dosya formatı işlenemedi: {file_extension}'
                }
        
        except Exception as e:
            logger.error(f"Error processing file: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """PDF dosyasını işle"""
        
        if not PDF_AVAILABLE:
            return {
                'success': False,
                'error': 'PDF işleme kütüphanesi yüklü değil. pip install pdfplumber'
            }
        
        try:
            # BytesIO kullanarak pdfplumber ile aç
            pdf_file = io.BytesIO(file_content)
            
            with pdfplumber.open(pdf_file) as pdf:
                text = ""
                page_count = len(pdf.pages)
                
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                # Metadata
                metadata = pdf.metadata or {}
                
                return {
                    'success': True,
                    'data': {
                        'filename': filename,
                        'text': text,
                        'page_count': page_count,
                        'file_type': 'pdf',
                        'metadata': {
                            'title': metadata.get('Title', ''),
                            'author': metadata.get('Author', ''),
                            'subject': metadata.get('Subject', '')
                        }
                    }
                }
        
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return {
                'success': False,
                'error': f'PDF işleme hatası: {str(e)}'
            }
    
    def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """DOCX dosyasını işle"""
        
        if not DOCX_AVAILABLE:
            return {
                'success': False,
                'error': 'DOCX işleme kütüphanesi yüklü değil. pip install python-docx'
            }
        
        try:
            # BytesIO kullanarak docx ile aç
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            # Tüm paragraflardan metin çıkar
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Tablolardan metin çıkar
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text += "\n" + row_text
            
            return {
                'success': True,
                'data': {
                    'filename': filename,
                    'text': text,
                    'page_count': 1,  # DOCX için sayfa sayısı belirsiz
                    'file_type': 'docx',
                    'metadata': {}
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return {
                'success': False,
                'error': f'DOCX işleme hatası: {str(e)}'
            }
    
    def _process_txt(self, path: Path) -> Dict[str, Any]:
        """TXT dosyasını işle (path'den)"""
        try:
            with open(path, 'r', encoding='utf-8') as f:
                text = f.read()
        except UnicodeDecodeError:
            with open(path, 'r', encoding='latin-1', errors='ignore') as f:
                text = f.read()
        
        return {
            'success': True,
            'data': {
                'filename': path.name,
                'text': text,
                'page_count': 1,  # TXT için sayfa sayısı 1
                'file_path': str(path),
                'file_size': path.stat().st_size
            }
        }
    
    def _process_txt_content(self, text: str, filename: str) -> Dict[str, Any]:
        """TXT içeriğini işle (string'den)"""
        return {
            'success': True,
            'data': {
                'filename': filename,
                'text': text,
                'page_count': 1,
                'file_path': filename,
                'file_size': len(text.encode('utf-8'))
            }
        }
    
    def _process_zip(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """ZIP dosyasını ayıkla ve içindeki dökümanları işle"""
        try:
            zip_file = io.BytesIO(file_content)
            extracted_texts = []
            extracted_files = []
            
            with zipfile.ZipFile(zip_file, 'r') as zip_ref:
                # ZIP içindeki tüm dosyaları listele
                file_list = zip_ref.namelist()
                
                for file_name in file_list:
                    # Güvenlik: ZIP içinde path traversal saldırılarına karşı kontrol
                    if '..' in file_name or file_name.startswith('/'):
                        logger.warning(f"[ZIP] Güvenlik: Şüpheli dosya adı atlandı: {file_name}")
                        continue
                    
                    # Sadece desteklenen formatları işle
                    file_ext = Path(file_name).suffix.lower().replace('.', '')
                    if file_ext not in ['pdf', 'docx', 'doc', 'txt', 'xls', 'xlsx']:
                        continue
                    
                    try:
                        # Dosyayı oku
                        file_data = zip_ref.read(file_name)
                        
                        # Dosyayı işle (recursive call)
                        if file_ext == 'pdf':
                            result = self._process_pdf(file_data, file_name)
                        elif file_ext in ['docx', 'doc']:
                            result = self._process_docx(file_data, file_name)
                        elif file_ext == 'txt':
                            try:
                                text_content = file_data.decode('utf-8')
                            except UnicodeDecodeError:
                                text_content = file_data.decode('latin-1', errors='ignore')
                            result = self._process_txt_content(text_content, file_name)
                        elif file_ext in ['xls', 'xlsx']:
                            result = self._process_excel(file_data, file_name)
                        else:
                            continue
                        
                        if result.get('success'):
                            extracted_texts.append(f"\n\n=== {file_name} ===\n\n{result['data'].get('text', '')}")
                            extracted_files.append(file_name)
                            logger.info(f"[ZIP] İşlendi: {file_name}")
                    
                    except Exception as e:
                        logger.warning(f"[ZIP] {file_name} işlenirken hata: {e}")
                        continue
            
            if not extracted_texts:
                return {
                    'success': False,
                    'error': 'ZIP dosyasında işlenebilir döküman bulunamadı'
                }
            
            combined_text = "\n".join(extracted_texts)
            
            return {
                'success': True,
                'data': {
                    'filename': filename,
                    'text': combined_text,
                    'page_count': len(extracted_files),
                    'file_type': 'zip',
                    'extracted_files': extracted_files,
                    'extracted_count': len(extracted_files),
                    'metadata': {
                        'source': 'zip_archive',
                        'files_in_zip': len(file_list)
                    }
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing ZIP: {e}")
            return {
                'success': False,
                'error': f'ZIP işleme hatası: {str(e)}'
            }
    
    def _process_excel(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Excel dosyasını işle (XLS/XLSX)"""
        if not EXCEL_AVAILABLE:
            return {
                'success': False,
                'error': 'Excel işleme kütüphanesi yüklü değil. pip install pandas openpyxl'
            }
        
        try:
            excel_file = io.BytesIO(file_content)
            
            # Tüm sheet'leri oku
            excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl' if filename.endswith('.xlsx') else None)
            
            text_parts = []
            sheet_count = 0
            
            for sheet_name, df in excel_data.items():
                sheet_count += 1
                # DataFrame'i string'e çevir
                sheet_text = f"\n\n=== Sheet: {sheet_name} ===\n\n"
                sheet_text += df.to_string(index=False)
                text_parts.append(sheet_text)
            
            combined_text = "\n".join(text_parts)
            
            return {
                'success': True,
                'data': {
                    'filename': filename,
                    'text': combined_text,
                    'page_count': sheet_count,
                    'file_type': 'excel',
                    'sheet_count': sheet_count,
                    'sheet_names': list(excel_data.keys()),
                    'metadata': {
                        'source': 'excel_file',
                        'total_rows': sum(len(df) for df in excel_data.values())
                    }
                }
            }
        
        except Exception as e:
            logger.error(f"Error processing Excel: {e}")
            return {
                'success': False,
                'error': f'Excel işleme hatası: {str(e)}'
            }
    
    def process_file_from_path(self, file_path: str) -> Dict[str, Any]:
        """Dosya yolundan işle"""
        
        try:
            path = Path(file_path)
            
            # TXT dosyası için özel işleme
            if path.suffix.lower() == '.txt':
                return self._process_txt(path)
            
            if not path.exists():
                return {
                    'success': False,
                    'error': f'Dosya bulunamadı: {file_path}'
                }
            
            file_extension = path.suffix.lower().replace('.', '')
            
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            if file_extension == 'pdf':
                return self._process_pdf(file_content, path.name)
            elif file_extension in ['docx', 'doc']:
                return self._process_docx(file_content, path.name)
            elif file_extension == 'txt':
                # TXT için path'den direkt oku (zaten _process_txt çağrıldı)
                return self._process_txt(path)
            elif file_extension == 'zip':
                return self._process_zip(file_content, path.name)
            elif file_extension in ['xls', 'xlsx']:
                return self._process_excel(file_content, path.name)
            else:
                return {
                    'success': False,
                    'error': f'Desteklenmeyen dosya formatı: {file_extension}'
                }
        
        except Exception as e:
            logger.error(f"Error processing file from path: {e}")
            return {
                'success': False,
                'error': str(e)
            }
    
    def download_document(self, url: str, filename: str) -> Optional[Path]:
        """
        Dokümanı indir ve yerel dosya yolunu döndür
        """
        try:
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            file_path = self.download_dir / filename
            
            with open(file_path, 'wb') as f:
                f.write(response.content)
            
            logger.info(f"✅ Doküman indirildi: {filename} ({len(response.content)} bytes)")
            return file_path
            
        except Exception as e:
            logger.error(f"❌ Doküman indirme hatası {url}: {e}")
            return None
    
    def mock_process_opportunity_documents(self, opportunity_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        [MOCK] Bir fırsat için tüm dokümanları işle - TEST AMAÇLIDIR
        Contains HARDCODED/FAKE data. Do not use in production analysis.
        """
        notice_id = opportunity_data.get('noticeId', 'unknown')
        logger.warning(f"⚠️ USING MOCK DATA for notice: {notice_id} - This is not real analysis!")
        
        # Mock document URLs (gerçek SAM.gov entegrasyonu için sam_lite_client.py kullanılacak)
        mock_documents = [
            {
                "name": f"{notice_id}_solicitation.pdf",
                "url": "https://example.com/mock.pdf",
                "type": "Solicitation"
            },
            {
                "name": f"{notice_id}_amendment.pdf", 
                "url": "https://example.com/mock2.pdf",
                "type": "Amendment"
            }
        ]

    def process_opportunity_documents(self, opportunity_data: Dict[str, Any]) -> Dict[str, Any]:
        """DEPRECATED: Use mock_process_opportunity_documents instead"""
        logger.warning("🚨 DEPRECATED FUNCTION CALLED: process_opportunity_documents returns MOCK DATA!")
        return self.mock_process_opportunity_documents(opportunity_data)
        
        processed_documents = []
        all_text = ""
        
        for doc in mock_documents:
            # Mock metin (gerçek implementasyonda download_document kullanılacak)
            mock_text = f"""
SOLICITATION DOCUMENT - {notice_id}

SECTION A - SOLICITATION/CONTRACT FORM

1. REQUIREMENTS:
- Hotel accommodation services for federal employees
- NAICS Code: 721110 - Hotels (except Casino Hotels) and Motels
- Period of Performance: 12 months
- Location: Washington DC metropolitan area

2. TECHNICAL REQUIREMENTS:
- Minimum 100 rooms available
- 24/7 front desk service
- Business center facilities
- Parking availability
- ADA compliance required

3. COMPLIANCE REQUIREMENTS:
- FAR 52.219-14 Limitations on Subcontracting
- Security clearance not required
- Past performance evaluation required
- Financial capability demonstration required

4. SUBMISSION REQUIREMENTS:
- Technical proposal (30 pages max)
- Price proposal (separate volume)
- Past performance references (minimum 3)
- Capability statements

5. EVALUATION CRITERIA:
- Technical Approach (40%)
- Past Performance (30%) 
- Price (30%)

6. KEY DATES:
- Questions Due: {opportunity_data.get('responseDeadline', 'TBD')}
- Proposal Due: {opportunity_data.get('responseDeadline', 'TBD')}
- Award Date: TBD
            """
            
            processed_doc = {
                "name": doc["name"],
                "type": doc["type"],
                "text_content": mock_text.strip(),
                "word_count": len(mock_text.split()),
                "processed_at": datetime.now().isoformat()
            }
            
            processed_documents.append(processed_doc)
            all_text += f"\n\n=== {doc['type']}: {doc['name']} ===\n\n{mock_text}"
        
        # Konsolide çıktı
        result = {
            "notice_id": notice_id,
            "processing_status": "completed",
            "processed_at": datetime.now().isoformat(),
            "documents": processed_documents,
            "consolidated_text": all_text.strip(),
            "total_word_count": len(all_text.split()),
            "document_count": len(processed_documents),
            "metadata": {
                "processor_version": "1.0",
                "extraction_method": "mock_data",  # gerçekte: "pdfplumber" veya "pypdf2"
                "supported_formats": self.supported_formats
            }
        }
        
        logger.info(f"✅ Doküman işleme tamamlandı: {notice_id} - {len(processed_documents)} doküman, {result['total_word_count']} kelime")
        return result


# Convenience function for external modules
def process_documents_for_opportunity(opportunity_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Bir fırsat için doküman işleme - dış modüllerden kullanım için
    """
    processor = DocumentProcessor()
    return processor.process_opportunity_documents(opportunity_data)

